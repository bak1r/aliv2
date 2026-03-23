"""Hukuki belge uretici — dilekce, savunma, itiraz sablonlari (DOCX)."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from tools.base import BaseTool
from core.config import BASE_DIR


OUTPUT_DIR = BASE_DIR / "data" / "documents"


class DocGeneratorTool(BaseTool):
    name = "belge_olustur"
    description = "Hukuki belge olusturur: dilekce, savunma, itiraz, temyiz. DOCX formatinda kaydeder."
    parameters = {
        "type": "object",
        "properties": {
            "belge_turu": {
                "type": "string",
                "description": "Belge turu: dilekce | savunma | itiraz | temyiz | genel"
            },
            "baslik": {"type": "string", "description": "Belge basligi"},
            "icerik": {"type": "string", "description": "Belge ana metni"},
            "mahkeme": {"type": "string", "description": "Mahkeme adi"},
            "dosya_no": {"type": "string", "description": "Dosya numarasi"},
            "davaci": {"type": "string", "description": "Davaci bilgisi"},
            "davali": {"type": "string", "description": "Davali bilgisi"},
        },
        "required": ["belge_turu", "icerik"],
    }

    def run(
        self,
        belge_turu: str = "genel",
        baslik: str = "",
        icerik: str = "",
        mahkeme: str = "",
        dosya_no: str = "",
        davaci: str = "",
        davali: str = "",
        **kw,
    ) -> str:
        if not icerik:
            return "Belge icerigi belirtilmedi."

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return "python-docx paketi bulunamadi. pip install python-docx"

        doc = Document()

        # Sayfa marjlari
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        tarih = datetime.now().strftime("%d.%m.%Y")

        if not baslik:
            baslik = {
                "dilekce": "DILEKCE",
                "savunma": "SAVUNMA BEYANI",
                "itiraz": "ITIRAZ DILEKCESI",
                "temyiz": "TEMYIZ DILEKCESI",
            }.get(belge_turu, "HUKUKI BELGE")

        # Tarih (sag ust)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(tarih)
        run.font.size = Pt(11)

        # Mahkeme
        if mahkeme:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(mahkeme.upper())
            run.bold = True
            run.font.size = Pt(13)

        # Baslik
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(baslik.upper())
        run.bold = True
        run.font.size = Pt(14)

        # Dosya bilgileri
        if dosya_no:
            doc.add_paragraph(f"Dosya No: {dosya_no}")

        if belge_turu in ("dilekce", "itiraz", "temyiz"):
            doc.add_paragraph("")
            if davaci:
                doc.add_paragraph(f"DAVACI    : {davaci}")
            if davali:
                doc.add_paragraph(f"DAVALI    : {davali}")
            doc.add_paragraph(f"KONU      : {baslik}")
            doc.add_paragraph("")

        # Ana icerik
        doc.add_paragraph("ACIKLAMALAR:")
        doc.add_paragraph("")

        for paragraph_text in icerik.split("\n"):
            p = doc.add_paragraph(paragraph_text)
            p.paragraph_format.space_after = Pt(6)

        # Sonuc ve talep
        doc.add_paragraph("")
        doc.add_paragraph("SONUC VE TALEP:")
        doc.add_paragraph("Yukarda aciklanan nedenlerle, gereginin yapilmasini saygilarimla arz ederim.")
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Imza
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Avukat: _______________\n")
        p.add_run("Imza  : _______________")

        # Alt not
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(f"Bu belge Ali v2 Avukat AI tarafindan {tarih} tarihinde olusturulmustur.")
        run.font.size = Pt(8)
        run.italic = True

        # Kaydet
        safe_name = "".join(c for c in baslik if c.isalnum() or c in " _-")[:40].strip() or belge_turu
        filename = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = OUTPUT_DIR / filename
        doc.save(str(filepath))

        # Dosyayi ac
        try:
            # platform import burada yapilir (circular import onleme)
            import sys
            if sys.platform == "win32":
                import os
                os.startfile(str(filepath))
            elif sys.platform == "darwin":
                import subprocess
                subprocess.Popen(["open", str(filepath)])
        except Exception:
            pass

        return f"Belge olusturuldu: {filepath}"

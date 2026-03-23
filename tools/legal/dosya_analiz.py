"""Toplu Dosya Analiz Araci — PDF, DOCX, TXT belgelerini okur, analiz eder, rapor uretir."""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from tools.base import BaseTool
from core.config import BASE_DIR


REPORT_DIR = BASE_DIR / "data" / "reports"
MAX_TOTAL_CHARS = 100_000

# Hukuki anahtar kelimeler
_SAVUNMA_KEYS = re.compile(
    r"\b(savunma|mudafi|itiraz|beraat|lehe|masumiyet|hakaniyet|orantililik|"
    r"meşru müdafaa|zorunluluk hali|haksız tahrik|indirim|cezanın ertelenmesi|"
    r"hükmün açıklanmasının geri bırakılması|hagb)\b", re.IGNORECASE
)
_IDDIA_KEYS = re.compile(
    r"\b(iddianame|savcılık|suçlama|isnat|mahkumiyet|aleyhe|cezalandırılma|"
    r"tutuklama|yakalama|suç|fail|sanık|müşteki)\b", re.IGNORECASE
)
_DELIL_KEYS = re.compile(
    r"\b(delil|belge|tutanak|rapor|bilirkişi|tanık ifadesi|keşif|iz|"
    r"kayıt|fotoğraf|video|ses kaydı|parmak izi|dna|otopsi|adli tıp|"
    r"ekspertiz|müzekkere|ek-|ekte)\b", re.IGNORECASE
)


def _extract_pdf(filepath: str) -> dict:
    """PDF dosyasından metin cikar."""
    try:
        import pdfplumber
    except ImportError:
        return {"filename": Path(filepath).name, "text": "", "page_count": 0,
                "word_count": 0, "error": "pdfplumber paketi bulunamadi"}

    pages_text = []
    try:
        with pdfplumber.open(filepath) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(text)
    except Exception as e:
        return {"filename": Path(filepath).name, "text": "", "page_count": 0,
                "word_count": 0, "error": str(e)}

    full_text = "\n".join(pages_text)
    return {
        "filename": Path(filepath).name,
        "text": full_text,
        "page_count": page_count,
        "word_count": len(full_text.split()),
    }


def _extract_docx(filepath: str) -> dict:
    """DOCX dosyasından metin cikar."""
    try:
        from docx import Document
    except ImportError:
        return {"filename": Path(filepath).name, "text": "", "page_count": 0,
                "word_count": 0, "error": "python-docx paketi bulunamadi"}

    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
    except Exception as e:
        return {"filename": Path(filepath).name, "text": "", "page_count": 0,
                "word_count": 0, "error": str(e)}

    return {
        "filename": Path(filepath).name,
        "text": full_text,
        "page_count": max(1, len(full_text) // 3000),  # tahmini
        "word_count": len(full_text.split()),
    }


def _extract_txt(filepath: str) -> dict:
    """TXT dosyasından metin cikar."""
    try:
        text = Path(filepath).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return {"filename": Path(filepath).name, "text": "", "page_count": 0,
                "word_count": 0, "error": str(e)}

    return {
        "filename": Path(filepath).name,
        "text": text,
        "page_count": 1,
        "word_count": len(text.split()),
    }


def _discover_files(folder: str) -> list[str]:
    """Klasordeki desteklenen dosyalari bul (recursive)."""
    supported = {".pdf", ".docx", ".doc", ".txt"}
    files = []
    for root, _dirs, filenames in os.walk(folder):
        for fname in sorted(filenames):
            if Path(fname).suffix.lower() in supported:
                files.append(os.path.join(root, fname))
    return files


def _extract_file(filepath: str) -> dict:
    """Dosya tipine gore metin cikar."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return _extract_docx(filepath)
    elif ext == ".txt":
        return _extract_txt(filepath)
    return {"filename": Path(filepath).name, "text": "", "page_count": 0,
            "word_count": 0, "error": "Desteklenmeyen format"}


def _truncate_texts(docs: list[dict], max_total: int = MAX_TOTAL_CHARS) -> bool:
    """Toplam metin boyutunu sinirla. Truncate yapildiysa True dondurur."""
    total = sum(len(d["text"]) for d in docs)
    if total <= max_total:
        return False

    per_file = max_total // max(len(docs), 1)
    for d in docs:
        if len(d["text"]) > per_file:
            d["text"] = d["text"][:per_file] + "\n... [kesildi]"
    return True


def _get_surrounding_context(text: str, match_start: int, match_end: int, context_chars: int = 200) -> str:
    """Esleme etrafindaki konteksti al."""
    start = max(0, match_start - context_chars)
    end = min(len(text), match_end + context_chars)
    prefix = "..." if start > 0 else ""
    suffix = "..." if end < len(text) else ""
    return prefix + text[start:end] + suffix


# ── Analiz fonksiyonlari ────────────────────────────────────────────

def _analiz_ozet(docs: list[dict]) -> tuple[str, list[str], list[str]]:
    """Dosyalarin ozetini cikar."""
    lines = ["## Dosya Ozeti\n"]
    for i, d in enumerate(docs, 1):
        preview = d["text"][:200].replace("\n", " ").strip()
        lines.append(f"**{i}. {d['filename']}**")
        lines.append(f"   Sayfa: {d['page_count']} | Kelime: {d['word_count']}")
        lines.append(f"   Onizleme: {preview}...")
        lines.append("")

    return "\n".join(lines), [], []


def _analiz_savunma(docs: list[dict]) -> tuple[str, list[str], list[str]]:
    """Savunma ve iddia unsurlarini analiz et."""
    lines = ["## Savunma Analizi\n"]
    bulgular = []
    oneriler = []

    for d in docs:
        savunma_hits = list(_SAVUNMA_KEYS.finditer(d["text"]))
        iddia_hits = list(_IDDIA_KEYS.finditer(d["text"]))

        lines.append(f"### {d['filename']}")

        if savunma_hits:
            lines.append(f"  Savunma unsurlari ({len(savunma_hits)} bulgu):")
            seen = set()
            for m in savunma_hits[:10]:
                kw = m.group().lower()
                if kw not in seen:
                    seen.add(kw)
                    ctx = _get_surrounding_context(d["text"], m.start(), m.end(), 150)
                    lines.append(f"  - **{m.group()}**: {ctx}")
            bulgular.append(f"{d['filename']}: {len(savunma_hits)} savunma unsuru bulundu")
        else:
            lines.append("  Savunma unsuru bulunamadi.")
            oneriler.append(f"{d['filename']}: Savunma argumanlari eksik — savunma stratejisi gelistirilmeli")

        if iddia_hits:
            lines.append(f"  Iddia unsurlari ({len(iddia_hits)} bulgu):")
            seen = set()
            for m in iddia_hits[:10]:
                kw = m.group().lower()
                if kw not in seen:
                    seen.add(kw)
                    ctx = _get_surrounding_context(d["text"], m.start(), m.end(), 150)
                    lines.append(f"  - **{m.group()}**: {ctx}")

        lines.append("")

    return "\n".join(lines), bulgular, oneriler


def _analiz_delil(docs: list[dict]) -> tuple[str, list[str], list[str]]:
    """Delil referanslarini cikar."""
    lines = ["## Delil Analizi\n"]
    bulgular = []
    oneriler = []
    all_evidence = {}

    for d in docs:
        hits = list(_DELIL_KEYS.finditer(d["text"]))
        lines.append(f"### {d['filename']}")

        if hits:
            seen = set()
            for m in hits[:15]:
                kw = m.group().lower()
                if kw not in seen:
                    seen.add(kw)
                    ctx = _get_surrounding_context(d["text"], m.start(), m.end(), 150)
                    lines.append(f"  - **{m.group()}**: {ctx}")
                    all_evidence.setdefault(kw, []).append(d["filename"])
            bulgular.append(f"{d['filename']}: {len(seen)} farkli delil turu referansi")
        else:
            lines.append("  Delil referansi bulunamadi.")
            oneriler.append(f"{d['filename']}: Delil referansi yok — dogrulanmali")

        lines.append("")

    # Capraz referans
    if all_evidence:
        lines.append("### Delil Capraz Referans Tablosu")
        for evidence, files in sorted(all_evidence.items()):
            lines.append(f"  - **{evidence}**: {', '.join(files)}")

    return "\n".join(lines), bulgular, oneriler


def _analiz_eksik(docs: list[dict]) -> tuple[str, list[str], list[str]]:
    """Eksikleri ve tutarsizliklari bul."""
    lines = ["## Eksik/Tutarsizlik Analizi\n"]
    bulgular = []
    oneriler = []

    # Tum dosyalardaki referanslari topla
    all_refs = {}
    for d in docs:
        for pattern in [_SAVUNMA_KEYS, _IDDIA_KEYS, _DELIL_KEYS]:
            for m in pattern.finditer(d["text"]):
                kw = m.group().lower()
                all_refs.setdefault(kw, set()).add(d["filename"])

    # Sadece tek dosyada gecen referanslar = potansiyel eksik
    single_refs = {k: v for k, v in all_refs.items() if len(v) == 1}
    multi_refs = {k: v for k, v in all_refs.items() if len(v) > 1}

    if single_refs:
        lines.append("### Sadece Tek Dosyada Gecen Kavramlar (Potansiyel Eksik)")
        for kw, files in sorted(single_refs.items()):
            f = list(files)[0]
            lines.append(f"  - **{kw}**: yalnizca *{f}* dosyasinda gecmektedir")
            bulgular.append(f"'{kw}' sadece {f} dosyasinda — diger belgelerde karsilanmamis olabilir")

    if multi_refs:
        lines.append("\n### Birden Fazla Dosyada Gecen Kavramlar (Tutarlilik)")
        for kw, files in sorted(multi_refs.items()):
            lines.append(f"  - **{kw}**: {', '.join(sorted(files))}")

    # Delil — tanik capraz kontrolu
    all_texts = " ".join(d["text"] for d in docs).lower()
    if "tanık" in all_texts and "tanık ifadesi" not in all_texts:
        oneriler.append("Tanik referansi var ancak tanik ifadesi belgesi bulunamadi")
    if "bilirkişi" in all_texts and "bilirkişi raporu" not in all_texts:
        oneriler.append("Bilirkisi referansi var ancak bilirkisi raporu bulunamadi")

    if not single_refs and not multi_refs:
        lines.append("Belirgin eksik veya tutarsizlik tespit edilemedi.")

    return "\n".join(lines), bulgular, oneriler


def _analiz_karsilastir(docs: list[dict]) -> tuple[str, list[str], list[str]]:
    """Belgeleri karsilastir."""
    lines = ["## Belge Karsilastirmasi\n"]
    bulgular = []

    # Temel istatistikler
    lines.append("### Temel Istatistikler")
    lines.append(f"{'Dosya':<40} {'Sayfa':>6} {'Kelime':>8}")
    lines.append("-" * 56)
    for d in docs:
        lines.append(f"{d['filename']:<40} {d['page_count']:>6} {d['word_count']:>8}")
    lines.append("")

    # Ortak ve farkli kavramlar
    doc_keywords = {}
    for d in docs:
        kws = set()
        for pattern in [_SAVUNMA_KEYS, _IDDIA_KEYS, _DELIL_KEYS]:
            kws.update(m.group().lower() for m in pattern.finditer(d["text"]))
        doc_keywords[d["filename"]] = kws

    if len(docs) >= 2:
        lines.append("### Kavram Karsilastirmasi")
        filenames = list(doc_keywords.keys())

        for i in range(len(filenames)):
            for j in range(i + 1, len(filenames)):
                f1, f2 = filenames[i], filenames[j]
                common = doc_keywords[f1] & doc_keywords[f2]
                only_f1 = doc_keywords[f1] - doc_keywords[f2]
                only_f2 = doc_keywords[f2] - doc_keywords[f1]

                lines.append(f"\n**{f1} vs {f2}:**")
                if common:
                    lines.append(f"  Ortak: {', '.join(sorted(common))}")
                if only_f1:
                    lines.append(f"  Sadece {f1}: {', '.join(sorted(only_f1))}")
                    bulgular.append(f"{f1} icinde olup {f2} icinde olmayan: {', '.join(sorted(only_f1))}")
                if only_f2:
                    lines.append(f"  Sadece {f2}: {', '.join(sorted(only_f2))}")

    return "\n".join(lines), bulgular, []


def _analiz_arama(docs: list[dict], terim: str) -> tuple[str, list[str], list[str]]:
    """Tum dosyalarda terim ara."""
    lines = [f"## Arama Sonuclari: '{terim}'\n"]
    bulgular = []

    if not terim:
        return "Arama terimi belirtilmedi.", [], ["arama_terimi parametresi gerekli"]

    pattern = re.compile(re.escape(terim), re.IGNORECASE)
    total_hits = 0

    for d in docs:
        hits = list(pattern.finditer(d["text"]))
        if hits:
            total_hits += len(hits)
            lines.append(f"### {d['filename']} ({len(hits)} sonuc)")
            for m in hits[:10]:
                ctx = _get_surrounding_context(d["text"], m.start(), m.end(), 200)
                lines.append(f"  - ...{ctx}...")
            if len(hits) > 10:
                lines.append(f"  ... ve {len(hits) - 10} sonuc daha")
            lines.append("")
            bulgular.append(f"{d['filename']}: {len(hits)} eslesme")
        else:
            lines.append(f"### {d['filename']}: sonuc yok")

    lines.insert(1, f"Toplam {total_hits} eslesme bulundu.\n")
    return "\n".join(lines), bulgular, []


def _generate_report(docs: list[dict], analiz_text: str, bulgular: list[str],
                     oneriler: list[str], gorev: str) -> str:
    """DOCX rapor uret ve kaydet."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "python-docx paketi bulunamadi."

    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    tarih = datetime.now().strftime("%d.%m.%Y %H:%M")
    gorev_labels = {
        "ozet": "Ozet Raporu", "savunma_analiz": "Savunma Analiz Raporu",
        "delil_analiz": "Delil Analiz Raporu", "eksik_bul": "Eksik Analiz Raporu",
        "karsilastir": "Karsilastirma Raporu", "arama": "Arama Raporu",
    }

    # Baslik
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOSYA ANALIZ RAPORU")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(gorev_labels.get(gorev, "Analiz Raporu"))
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph(f"Tarih: {tarih}")
    doc.add_paragraph(f"Analiz edilen dosya sayisi: {len(docs)}")
    doc.add_paragraph(f"Toplam kelime sayisi: {sum(d['word_count'] for d in docs)}")
    doc.add_paragraph("")

    # Dosya listesi
    p = doc.add_paragraph()
    run = p.add_run("DOSYA LISTESI")
    run.bold = True
    run.font.size = Pt(12)

    for i, d in enumerate(docs, 1):
        doc.add_paragraph(f"{i}. {d['filename']} — {d['page_count']} sayfa, {d['word_count']} kelime")

    doc.add_paragraph("")

    # Analiz
    p = doc.add_paragraph()
    run = p.add_run("ANALIZ")
    run.bold = True
    run.font.size = Pt(12)

    # Markdown-ish isaretleri temizle
    clean_text = analiz_text.replace("##", "").replace("**", "").replace("*", "")
    for line in clean_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph("")

    # Bulgular
    if bulgular:
        p = doc.add_paragraph()
        run = p.add_run("BULGULAR")
        run.bold = True
        run.font.size = Pt(12)

        for b in bulgular:
            doc.add_paragraph(f"- {b}")

    # Oneriler
    if oneriler:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("ONERILER")
        run.bold = True
        run.font.size = Pt(12)

        for o in oneriler:
            doc.add_paragraph(f"- {o}")

    # Alt not
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Bu rapor Ali v2 tarafindan {tarih} tarihinde otomatik olusturulmustur.")
    run.font.size = Pt(8)
    run.italic = True

    # Kaydet
    filename = f"analiz_{gorev}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = REPORT_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


# ── Ana Arac Sinifi ──────────────────────────────────────────────────

class DosyaAnalizTool(BaseTool):
    name = "dosya_analiz"
    description = (
        "Toplu dosya analizi — PDF ve Word belgelerini okur, analiz eder, rapor uretir. "
        "Gorevler: ozet, savunma_analiz, delil_analiz, eksik_bul, karsilastir, arama."
    )
    parameters = {
        "type": "object",
        "properties": {
            "klasor": {
                "type": "string",
                "description": "Analiz edilecek dosyalarin bulundugu klasor yolu",
            },
            "gorev": {
                "type": "string",
                "description": (
                    "Analiz gorevi: ozet | savunma_analiz | delil_analiz | "
                    "eksik_bul | karsilastir | arama"
                ),
                "enum": ["ozet", "savunma_analiz", "delil_analiz", "eksik_bul", "karsilastir", "arama"],
            },
            "arama_terimi": {
                "type": "string",
                "description": "Arama terimi (gorev 'arama' oldugunda kullanilir)",
            },
            "cikti_format": {
                "type": "string",
                "description": "Cikti formati: ozet (metin) veya rapor (DOCX dosyasi)",
                "enum": ["ozet", "rapor"],
                "default": "ozet",
            },
        },
        "required": ["klasor", "gorev"],
    }

    def run(
        self,
        klasor: str = "",
        gorev: str = "ozet",
        arama_terimi: str = "",
        cikti_format: str = "ozet",
        **kw,
    ) -> str:
        if not klasor:
            return "Klasor yolu belirtilmedi."

        folder = Path(klasor).expanduser()
        if not folder.exists():
            return f"Klasor bulunamadi: {folder}"
        if not folder.is_dir():
            return f"Belirtilen yol bir klasor degil: {folder}"

        # Dosyalari bul
        files = _discover_files(str(folder))
        if not files:
            return f"Klasorde desteklenen dosya bulunamadi (.pdf, .docx, .doc, .txt): {folder}"

        # Metinleri cikar
        docs = []
        errors = []
        for f in files:
            result = _extract_file(f)
            if "error" in result:
                errors.append(f"{result['filename']}: {result['error']}")
            if result["text"]:
                docs.append(result)

        if not docs:
            error_msg = "Hicbir dosyadan metin cikarilamadi."
            if errors:
                error_msg += "\nHatalar:\n" + "\n".join(errors)
            return error_msg

        # Truncation kontrolu
        truncated = _truncate_texts(docs)

        # Analiz yap
        gorev_map = {
            "ozet": lambda: _analiz_ozet(docs),
            "savunma_analiz": lambda: _analiz_savunma(docs),
            "delil_analiz": lambda: _analiz_delil(docs),
            "eksik_bul": lambda: _analiz_eksik(docs),
            "karsilastir": lambda: _analiz_karsilastir(docs),
            "arama": lambda: _analiz_arama(docs, arama_terimi),
        }

        if gorev not in gorev_map:
            return f"Gecersiz gorev: {gorev}. Secenekler: {', '.join(gorev_map.keys())}"

        analiz_text, bulgular, oneriler = gorev_map[gorev]()

        # Sonuc olustur
        toplam_kelime = sum(d["word_count"] for d in docs)
        dosya_bilgileri = [
            {"filename": d["filename"], "page_count": d["page_count"], "word_count": d["word_count"]}
            for d in docs
        ]

        result_parts = [
            f"# Dosya Analiz Sonucu",
            f"Dosya sayisi: {len(docs)} | Toplam kelime: {toplam_kelime}",
        ]

        if truncated:
            result_parts.append("⚠ Toplam metin boyutu 100K karakteri astigindan dosyalar kesildi.")

        if errors:
            result_parts.append(f"\nOkunamayan dosyalar: {', '.join(errors)}")

        result_parts.append("")
        result_parts.append(analiz_text)

        if bulgular:
            result_parts.append("\n### Temel Bulgular")
            for b in bulgular:
                result_parts.append(f"- {b}")

        if oneriler:
            result_parts.append("\n### Oneriler")
            for o in oneriler:
                result_parts.append(f"- {o}")

        # DOCX rapor
        if cikti_format == "rapor":
            report_path = _generate_report(docs, analiz_text, bulgular, oneriler, gorev)
            if report_path.endswith(".docx"):
                result_parts.append(f"\nRapor olusturuldu: {report_path}")
            else:
                result_parts.append(f"\nRapor olusturulamadi: {report_path}")

        return "\n".join(result_parts)
